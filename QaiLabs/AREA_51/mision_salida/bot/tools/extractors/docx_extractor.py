"""
Extractor de texto de archivos DOCX (Word).

Dependencias: pip install python-docx
"""

from docx import Document
from pathlib import Path


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extrae texto de un archivo DOCX.
    
    Args:
        docx_path: Ruta al archivo DOCX
        
    Returns:
        Texto extraído del documento
    """
    try:
        doc = Document(docx_path)
        text_parts = []
        
        # Extraer texto de todos los párrafos
        for para in doc.paragraphs:
            if para.text.strip():  # Solo agregar párrafos no vacíos
                text_parts.append(para.text)
        
        # Extraer texto de tablas (si existen)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
        
    except Exception as e:
        return f"❌ Error extrayendo texto del DOCX: {str(e)}"


# Ejemplo de uso
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Uso: python docx.py <ruta_archivo.docx>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if not Path(docx_file).exists():
        print(f"❌ Archivo no encontrado: {docx_file}")
        sys.exit(1)
    
    print(f"📄 Extrayendo texto de: {docx_file}")
    text = extract_text_from_docx(docx_file)
    
    print("\n" + "="*50)
    print(text[:500] + "..." if len(text) > 500 else text)
    print("="*50)
    print(f"\n✅ Total extraído: {len(text)} caracteres")
